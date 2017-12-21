from docx import Document
from docx.shared import Inches
from datetime import date, timedelta

from docx.styles.latent import LatentStyles
d1 = date(2018, 1, 1)  # start date
d2 = date(2018, 5, 11)  # end date

delta = d2 - d1         # timedelta

document = Document()
tot = delta.days + 1
maanden = ['Januari', 'Februari', 'Maart', 'April', 'Mei']
for i in range(tot):
    day = (d1 + timedelta(days=i))


    document.add_heading('Countdown to B-Day...', 0)

    document.add_paragraph('nog', style='Heading 4')
    p = document.add_paragraph(str(tot -i-1), style='IntenseQuote')
    document.add_paragraph('dagen', style='Heading 5')
    document.add_paragraph('%d %s %d' % (day.day, maanden[day.month -1], day.year), style='Heading 6')

    document.add_page_break()

#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True

#document.add_heading('Heading, level 1', level=1)
#document.add_paragraph('Intense quote', style='IntenseQuote')

#document.add_paragraph(
#    'first item in unordered list', style='ListBullet'
#)
#document.add_paragraph(
#    'first item in ordered list', style='ListNumber'
#)

#document.add_picture('monty-truth.png', width=Inches(1.25))

#table = document.add_table(rows=1, cols=3)
#hdr_cells = table.rows[0].cells
#hdr_cells[0].text = 'Qty'
#hdr_cells[1].text = 'Id'
#hdr_cells[2].text = 'Desc'
#for item in ['jan','feb','maa','april','mei']:
#    row_cells = table.add_row().cells
#   row_cells[0].text = str(item[0])
#    row_cells[1].text = str(item[1])
#    row_cells[2].text = item[2]

document.add_page_break()

document.save('demo.docx')
